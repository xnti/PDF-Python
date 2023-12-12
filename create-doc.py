from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Test Document', 0)

# Add some content to the document
doc.add_paragraph("This is a test document for the DOCX to PDF conversion.")
doc.add_paragraph("It contains multiple paragraphs to demonstrate the conversion process.")

# Save the document
file_path = 'test_document.docx'
doc.save(file_path)

file_path
